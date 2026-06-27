#!/usr/bin/env python3
"""Заполнение официальных форм конкурса Белинфонда для МЦ «Кравира».

Источник шаблонов: https://konkurs.belinfund.by/participants
Выход: docs/konkurs/*.docx

Перед подачей замените поля [УНП], [ФИО директора], [адрес], [телефон], [email].
"""
from __future__ import annotations

import re
import sys
import urllib.parse
import urllib.request
from pathlib import Path

from docx import Document
from docx.shared import Pt

_SCRIPTS = Path(__file__).resolve().parent
if str(_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS))

from konkurs_bp_sections import RESUME, SECTIONS, TITLE_PAGE_LINES, TOC
from konkurs_docx_helpers import (
    dash,
    enrich_business_plan,
    format_title_page,
    generate_charts,
    normalize_document_dashes,
)
from konkurs_org import (
    CONTACT_PERSON,
    DIRECTOR,
    EMAIL,
    NOMINATION,
    ORG_ADDRESS,
    ORG_NAME,
    ORG_SHORT,
    ORG_UNP,
    PASSPORT_IP,
    PHONE,
    PROJECT_NAME,
    TEAM,
    WEB,
)

ROOT = Path(__file__).resolve().parents[1]
ASSETS_DIR = ROOT / "docs" / "konkurs" / "_assets"
OUT_DIR = ROOT / "docs" / "konkurs"
TMP = OUT_DIR / "_templates"

BASE_URL = "https://konkurs.belinfund.by/storage/docs/"
TEMPLATES = {
    "zayavka": "Заявка форма_1.docx",
    "passport": "Паспорт форма_1.docx",
    "business": "Бизнес план форма_1.docx",
    "strategy": "Стратегия форма_1.docx",
}

PASSPORT_DESC = """\
В Республике Беларусь действует порядка 478 официальных клинических протоколов Минздрава; частная амбулатория обязана вести консультативные заключения (КЗ) по стандартам и передавать данные в ЦИСЗ (FHIR BY). Врач не может держать в памяти все протоколы; методслужба не проверяет каждое КЗ до подписи ЭЦП. Пациенты не имеют прозрачного инструмента оценки качества оказанной помощи по их заключению.

Protocol - отечественный продукт МЦ «Кравира»: (1) подбор клинических протоколов по МКБ-10; (2) детерминированная оценка КЗ (8 блоков, 482+ правил, цитаты из PDF Минздрава); (3) готовность к ЦИСЗ; (4) решение send_gate для МИС. LLM - только пояснения; итог для блокировки - по правилам.

Три канала монетизации: B2B - от 0,99 BYN за проверку КЗ в клинике (пакеты для потока ~25 000 КЗ/мес в Кравире); B2B API в МИС «Айболит»; B2C - для физических лиц: загрузка своего КЗ и отчёт «соответствует ли помощь клиническим протоколам» (от 4,99 BYN). Рабочий прототип развёрнут. Не заменяет врача, МЭЭ и валидацию ЦИСЗ."""

PASSPORT_CONSUMERS = """\
B2B: частные ОЗ и сети клиник РБ, методслужбы, разработчики МИС (Айболит/EPAM). Якорь - МЦ «Кравира» (~25 000 КЗ/мес, ~1% рынка частных ОЗ). B2C: физические лица - пациенты, проверяющие своё КЗ и качество оказанной услуги. Государственный сектор - перспектива после пилота."""

PASSPORT_ADVANTAGES = """\
1) Корпус ~478 КП Минздрава с цитатами. 2) Детерминированный scoring и send_gate. 3) Клиника + ЦИСЗ в одном контуре. 4) L0 <2 с для МИС. 5) Три канала дохода: клиники, МИС, физлица. 6) Пилот на потоке Кравиры. 7) On-prem в РБ, защита ПДн."""

PASSPORT_TIMELINE = """\
2025-Q2 2026: пилот L0 в Кравире. Q3 2026-Q1 2027: API в МИС Айболит. 2027: B2B 3-5 ОЗ, запуск B2C-витрины для физлиц. 2028-2029: 5-15% рынка частных ОЗ, масштабирование B2C."""

PASSPORT_PRODUCT_CERT = (
    "Услуга и ПО Protocol: предимпортная экспертиза КЗ для клиник (B2B) "
    "и сервис самопроверки КЗ для физических лиц (B2C) "
    "с отчётом по клиническим протоколам Минздрава РБ."
)

PASSPORT_ACHIEVEMENTS = """\
Веб-сервис и API (FastAPI); корпус ~478 протоколов Минздрава, 482+ правил; tiering L0/L1/L2; чек-лист ЦИСЗ 3.2.1; документация architecture-kravira-fhir-mis.pdf. Пилот на площадке МЦ «Кравира» (приказ директора о пилоте - оформляется до 01.08.2026). Переговоры с EPAM о интеграции L0 API в МИС «Айболит» (письмо о намерениях - до 01.08.2026)."""

PASSPORT_EXTRA = """\
Приложения: architecture-kravira-fhir-mis.pdf, mvp-presentation.html, pre-sign-checklist. Демо: веб-интерфейс проверки КЗ."""

STRATEGY_YEAR = "2027"
STRATEGY_LEVELS = (
    "отмечены: определены способы монетизации (B2B микроплатёж + B2C для физлиц), "
    "ценовая политика, каналы продаж (МИС, клиники, веб-витрина); "
    "определён потенциальный заказчик - МЦ «Кравира»; "
    "предварительный вывод на рынок - письма заинтересованности оформляются до 01.08.2026."
)
STRATEGY_METHODS = (
    "реализация услуг (B2B проверка КЗ, B2C самопроверка для физлиц); "
    "предоставление права использования (лицензия API для МИС); "
    "заинтересованность (письма клиник и интегратора МИС)."
)
STRATEGY_YEAR_PLAN = """\
1) Завершить пилот L0 в Кравире и метрики B2B.
2) Масштабировать B2C (patient.html развёрнут): ERIP, tier 4,99-14,99 BYN, QR/SMS rev-share, политика ПДн.
3) Интегрировать API в МИС «Айболит».
4) QR-кампания «Проверьте своё заключение» в Кравире и партнёрских ОЗ.
5) Подключить 2 внешние клиники B2B."""
STRATEGY_5Y = """\
2026-2027: Кравира B2B + B2C MVP (patient.html) + ERIP. 2028: 5-10 ОЗ B2B, 30-50k B2C проверок/год. 2029: white-label МИС, 8% рынка B2B. 2030: госсектор/сети. 2031: регион ЕАЭС, адаптация корпуса КП."""


def _download_templates() -> None:
    TMP.mkdir(parents=True, exist_ok=True)
    import ssl

    ctx = ssl.create_default_context()
    for fname in TEMPLATES.values():
        dest = TMP / fname
        if dest.is_file() and dest.stat().st_size > 1000:
            continue
        url = BASE_URL + urllib.parse.quote(fname)
        print(f"Download {fname}")
        try:
            urllib.request.urlretrieve(url, dest, context=ctx)
        except ssl.SSLError:
            # fallback для сред без цепочки CA
            urllib.request.urlretrieve(url, dest, context=ssl._create_unverified_context())


def _set_cell(table, row: int, col: int, text: str) -> None:
    table.rows[row].cells[col].text = dash(text)


def _fill_zayavka(src: Path, dest: Path) -> None:
    doc = Document(str(src))
    t = doc.tables[0]
    _set_cell(t, 0, 1, NOMINATION)
    _set_cell(t, 1, 1, PROJECT_NAME)
    # Юридическое лицо
    _set_cell(t, 14, 1, ORG_NAME)
    _set_cell(t, 15, 1, DIRECTOR)
    _set_cell(t, 16, 1, ORG_ADDRESS)
    _set_cell(t, 17, 1, ORG_UNP)
    _set_cell(t, 18, 1, CONTACT_PERSON)
    _set_cell(t, 19, 1, PHONE)
    if len(t.rows) > 20:
        _set_cell(t, 20, 1, EMAIL)
    if len(t.rows) > 21:
        _set_cell(t, 21, 1, TEAM)
    if len(t.rows) > 22:
        _set_cell(t, 22, 1, WEB)
    normalize_document_dashes(doc)
    doc.save(str(dest))


def _fill_passport(src: Path, dest: Path) -> None:
    doc = Document(str(src))
    t = doc.tables[0]
    _set_cell(t, 0, 2, PROJECT_NAME)
    _set_cell(t, 1, 2, PASSPORT_DESC.strip())
    directions = t.rows[2].cells[2].text
    selected = "☑ Медицинские науки и технологии\n☑ Информатика, информатизация и космические исследования"
    _set_cell(t, 2, 2, selected + "\n\n(исходный перечень формы: " + directions[:200] + "…)")
    novelty = "☑ Нет аналогов в стране, есть за рубежом"
    _set_cell(t, 3, 2, novelty)
    stage = "☑ Работающий прототип"
    _set_cell(t, 4, 2, stage)
    _set_cell(t, 5, 2, PASSPORT_CONSUMERS.strip())
    _set_cell(t, 6, 2, PASSPORT_ADVANTAGES.strip())
    _set_cell(
        t,
        7,
        2,
        "☑ Используются либо планируются к использованию объекты интеллектуальной "
        "собственности, права на которые подтверждаются соответствующими документами\n\n"
        "Пояснения:\n" + PASSPORT_IP.strip(),
    )
    _set_cell(t, 8, 2, PASSPORT_TIMELINE.strip())
    _set_cell(t, 9, 2, f"☑ Согласен\n\n{PASSPORT_PRODUCT_CERT}")
    _set_cell(t, 10, 2, PASSPORT_ACHIEVEMENTS.strip())
    _set_cell(t, 11, 2, PASSPORT_EXTRA.strip())
    normalize_document_dashes(doc)
    doc.save(str(dest))


def _insert_paragraph_after(paragraph, text: str):
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.add_run(dash(text))
    return new_para


def _replace_underscore_block(doc: Document, header_prefix: str, content: str) -> None:
    """Вставить текст раздела: убрать линию подчёркиваний в том же или следующем абзаце."""
    body = dash(content.strip())
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if not txt.startswith(header_prefix):
            continue
        if re.search(r"_{5,}", txt):
            header_only = re.sub(r"_{5,}.*", "", txt).strip()
            p.clear()
            p.add_run(header_only)
            _insert_paragraph_after(p, body)
            return
        if i + 1 < len(doc.paragraphs):
            nxt = doc.paragraphs[i + 1].text.strip()
            if re.match(r"^_{5,}", nxt):
                para = doc.paragraphs[i + 1]
                para.clear()
                para.add_run(body)
                return
        _insert_paragraph_after(p, body)
        return
    doc.add_paragraph(body)


def _fill_business(src: Path, dest: Path) -> None:
    doc = Document(str(src))
    # titul
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("1. Титульный лист"):
            if i + 1 < len(doc.paragraphs) and re.match(r"^_{5,}", doc.paragraphs[i + 1].text.strip()):
                format_title_page(doc.paragraphs[i + 1], TITLE_PAGE_LINES)
            break
    for p in doc.paragraphs:
        if p.text.strip().startswith("2. Содержание"):
            p.clear()
            p.add_run(dash("2. Содержание\n\n" + TOC.strip()))
            break
    _replace_underscore_block(doc, "3. Резюме", RESUME)
    for prefix, body in SECTIONS.items():
        _replace_underscore_block(doc, prefix, body)
    charts = generate_charts(ASSETS_DIR)
    enrich_business_plan(doc, charts)
    normalize_document_dashes(doc)
    doc.save(str(dest))


def _fill_strategy(src: Path, dest: Path) -> None:
    doc = Document(str(src))
    replacements = {
        3: PROJECT_NAME,
        6: NOMINATION,
        9: f"{ORG_NAME} ({ORG_SHORT})",
        15: f"опытный образец, B2B-пилот в Кравире, B2C MVP patient.html (масштаб ERIP и rev-share - {STRATEGY_YEAR}).",
    }
    for idx, text in replacements.items():
        if idx < len(doc.paragraphs):
            p = doc.paragraphs[idx]
            if re.match(r"^_{5,}", p.text.strip()) or idx in (3, 6, 9, 15):
                p.clear()
                p.add_run(text)
    # уровень коммерциализации - вставка после списка
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Выбор способа коммерциализации"):
            # next content paragraph with underscores
            for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
                if re.match(r"^_{5,}", doc.paragraphs[j].text.strip()):
                    doc.paragraphs[j].clear()
                    doc.paragraphs[j].add_run(STRATEGY_METHODS)
                    break
            break
    for i, p in enumerate(doc.paragraphs):
        if "Описание стратегии коммерциализации (план действий) на ближайший год" in p.text:
            if i + 1 < len(doc.paragraphs):
                para = doc.paragraphs[i + 1]
                if re.match(r"^_{5,}", para.text.strip()):
                    para.clear()
                    para.add_run(STRATEGY_YEAR_PLAN.strip())
            break
    for i, p in enumerate(doc.paragraphs):
        if "Стратегия коммерциализации на последующие 5 лет" in p.text:
            if i + 1 < len(doc.paragraphs):
                para = doc.paragraphs[i + 1]
                if re.match(r"^_{5,}", para.text.strip()):
                    para.clear()
                    para.add_run(STRATEGY_5Y.strip())
            break
    # note on commercialization level after header block
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Уровень коммерциализации на момент подачи заявки:":
            ins = p.insert_paragraph_before(STRATEGY_LEVELS)
            ins.runs[0].italic = True
            break
    normalize_document_dashes(doc)
    doc.save(str(dest))


def main() -> None:
    _download_templates()
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    _fill_zayavka(TMP / TEMPLATES["zayavka"], OUT_DIR / "01_Zayavka_Kravira_Protocol.docx")
    _fill_passport(TMP / TEMPLATES["passport"], OUT_DIR / "02_Pasport_Kravira_Protocol.docx")
    _fill_business(TMP / TEMPLATES["business"], OUT_DIR / "03_Biznes_plan_Kravira_Protocol.docx")
    _fill_strategy(TMP / TEMPLATES["strategy"], OUT_DIR / "04_Strategiya_Kravira_Protocol.docx")
    print("OK:", OUT_DIR)


if __name__ == "__main__":
    main()
