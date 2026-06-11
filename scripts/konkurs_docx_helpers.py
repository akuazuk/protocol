"""Таблицы, графики и приложения для бизнес-плана конкурса Белинфонд."""
from __future__ import annotations

import re
import sys
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

_SCRIPTS = Path(__file__).resolve().parent
if str(_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_SCRIPTS))

from konkurs_finance import (  # noqa: E402
    B2C_AVG_PRICE,
    B2C_TIERS,
    B2C_TAM_TOUCHES_YEAR,
    CERTIFICATE_BYN,
    CLINIC_B2C_REVSHARE,
    FIN_Y1,
    FIN_Y2,
    FIN_Y3,
    KRAVIRA_B2B_YEAR,
    KRAVIRA_KZ_MONTH,
    MARKET_KZ_MONTH,
    MARKET_KZ_YEAR,
    ROI_METHODIST_SAVING,
    ROI_NET,
    ROI_PROTOCOL_COST,
    ROI_TOTAL_SAVING,
    SAM_KZ_YEAR,
    SOM_Y3_KZ_YEAR,
    TAM_B2B_CEILING_YEAR_K,
    TAM_REVENUE_YEAR,
    Y2_MARKET_SHARE,
    Y3_MARKET_SHARE,
    clinic_revshare_byn,
    ebitda_k,
    ebitda_month_k,
    total_rev_k,
)
from konkurs_scenarios import (  # noqa: E402
    ALL_SCENARIOS_Y3,
    B2C_PROTOCOL_PER_CHECK,
    CHANNEL_OUTLOOK,
    FIN_Y3_SYNC,
    MONTHLY_Y3_CAUTIOUS,
    PENETRATION_SENSITIVITY,
    SCENARIO_CAUTIOUS,
    SCENARIO_BASE,
    SCENARIO_OPTIMISTIC,
    TAM_BRIDGE,
    TAM_CEILING_B2B_YEAR_K,
    b2c_protocol_k,
)
from konkurs_chart_style import BG_FIG, COLOR_NEGATIVE, COLORS, apply_rc, style_ax  # noqa: E402
from konkurs_market import (  # noqa: E402
    CISZ_DRIVERS,
    COMPETITOR_MATRIX,
    INVESTMENT_PLAN,
    OPEX_BREAKDOWN,
    RB_MARKET_TABLE,
)

GREEN = "126B5C"
GREEN_LIGHT = "E8F5F1"
AMBER = "D97706"
OCEAN = "0C6B94"


def dash(text: str) -> str:
    """Короткое тире вместо длинного/среднего."""
    return text.replace("\u2014", "-").replace("\u2013", "-")


def normalize_document_dashes(doc: Document) -> None:
    """Заменить длинные тире во всём документе (включая шаблон формы)."""
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                run.text = dash(run.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = dash(run.text)


def _shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(dash(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def add_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    *,
    bold: bool = False,
    italic: bool = False,
    size: int = 11,
    space_before: int = 6,
    space_after: int = 6,
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(dash(text))
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
    fmt = new_para.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    return new_para


def add_table_after(
    paragraph: Paragraph,
    headers: Sequence[str],
    rows: Sequence[Sequence[str]],
    *,
    caption: str | None = None,
) -> Paragraph:
    anchor = paragraph
    if caption:
        anchor = add_paragraph_after(anchor, caption, bold=True, size=10, space_after=3)

    doc = paragraph.part.document
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    anchor._p.addnext(table._tbl)

    for col, header in enumerate(headers):
        cell = table.rows[0].cells[col]
        _shade_cell(cell, GREEN)
        _set_cell_text(cell, header, bold=True, color=RGBColor(255, 255, 255), size=9)

    for r_idx, row in enumerate(rows, start=1):
        fill = GREEN_LIGHT if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, value in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            _shade_cell(cell, fill)
            _set_cell_text(cell, str(value), size=9)

    tail = OxmlElement("w:p")
    table._tbl.addnext(tail)
    return Paragraph(tail, paragraph._parent)


def add_picture_after(paragraph: Paragraph, image_path: Path, width_inches: float = 5.8) -> Paragraph:
    doc = paragraph.part.document
    cap = add_paragraph_after(paragraph, "", space_before=8, space_after=2)
    pic_p = OxmlElement("w:p")
    cap._p.addnext(pic_p)
    pic_para = Paragraph(pic_p, paragraph._parent)
    pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_para.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    tail = add_paragraph_after(pic_para, "", space_before=4, space_after=8)
    return tail


def format_title_page(
    paragraph: Paragraph,
    lines: Sequence[tuple[str, bool, int, str]],
) -> None:
    """Оформить титульный лист: центрирование заголовков, размеры шрифта."""
    paragraph.clear()
    anchor = paragraph
    for idx, (text, bold, size, align) in enumerate(lines):
        if idx == 0:
            p = anchor
        else:
            p = add_paragraph_after(anchor, "", space_before=2, space_after=2)
            anchor = p
        if not text:
            continue
        run = p.add_run(dash(text))
        run.bold = bold
        run.font.size = Pt(size)
        if align == "center":
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif align == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def find_section_body(doc: Document, section_num: str) -> Paragraph | None:
    """Абзац с текстом раздела сразу после заголовка «N. …»."""
    want = re.compile(rf"^{re.escape(section_num)}\s*\.")
    after_header = False
    for p in doc.paragraphs:
        t = p.text.strip()
        if want.match(t):
            after_header = True
            continue
        if after_header:
            if t and not re.match(r"^\d+\s*\.", t):
                return p
            if re.match(r"^\d+\s*\.", t):
                return None
    return None


def generate_charts(assets_dir: Path) -> dict[str, Path]:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    assets_dir.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    colors = COLORS
    apply_rc(plt)

    def _save(fig, path: Path) -> None:
        fig.tight_layout()
        fig.savefig(path, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
        plt.close(fig)

    # Выручка по годам (B2B + B2C)
    fig, ax = plt.subplots(figsize=(7, 3.8))
    years = ["2027", "2028", "2029"]
    b2b = [FIN_Y1["b2b_k"], FIN_Y2["b2b_k"], FIN_Y3["b2b_k"]]
    b2c = [FIN_Y1["b2c_k"], FIN_Y2["b2c_k"], FIN_Y3["b2c_k"]]
    x = range(len(years))
    w = 0.35
    ax.bar([i - w / 2 for i in x], b2b, width=w, label="B2B, тыс. BYN", color=colors[0])
    ax.bar([i + w / 2 for i in x], b2c, width=w, label="B2C, тыс. BYN", color=colors[4])
    ax.set_xticks(list(x))
    ax.set_xticklabels(years)
    ax.set_ylabel("тыс. BYN / год")
    ax.set_title(f"Прогноз выручки Protocol · осторожный сценарий (SOM {Y3_MARKET_SHARE:.0%} TAM)")
    ax.legend(loc="upper left", fontsize=8)
    style_ax(ax)
    p1 = assets_dir / "chart_revenue.png"
    _save(fig, p1)
    paths["revenue"] = p1

    # TAM / SAM / SOM
    fig, ax = plt.subplots(figsize=(7, 3.8))
    labels = [
        f"TAM\n{MARKET_KZ_YEAR // 1_000_000} млн КЗ/год",
        f"SAM 5%\n{SAM_KZ_YEAR // 1_000_000:,} млн КЗ".replace(",", " "),
        f"SOM год 3\n{SOM_Y3_KZ_YEAR // 1_000_000:,} млн КЗ".replace(",", " "),
    ]
    values = [MARKET_KZ_YEAR / 1_000_000, SAM_KZ_YEAR / 1_000_000, SOM_Y3_KZ_YEAR / 1_000_000]
    bars = ax.bar(labels, values, color=[colors[2], colors[1], colors[0]])
    ax.set_ylabel("млн КЗ / год")
    ax.set_title(f"Рынок частных ОЗ РБ: TAM 30 млн → SAM → SOM 8% = {SOM_Y3_KZ_YEAR // 1_000_000} млн КЗ/год")
    for bar, val in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.3, f"{val}", ha="center", fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p2 = assets_dir / "chart_market.png"
    fig.savefig(p2, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["market"] = p2

    # Тарифные пакеты B2B
    fig, ax = plt.subplots(figsize=(7, 3.2))
    tiers = ["Старт\nдо 1k", "Клиника\n10k", "Сеть\n25k+"]
    prices = [0.99, 0.79, 0.69]
    bars = ax.barh(tiers, prices, color=colors[:3])
    ax.set_xlabel("BYN за проверку L0")
    ax.set_title("Тарифная лестница B2B (микроплатёж за КЗ)")
    ax.set_xlim(0, 1.1)
    for bar, val in zip(bars, prices):
        ax.text(val + 0.02, bar.get_y() + bar.get_height() / 2, f"{val:.2f}", va="center", fontsize=9)
    ax.grid(axis="x", alpha=0.3)
    fig.tight_layout()
    p3 = assets_dir / "chart_pricing.png"
    fig.savefig(p3, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["pricing"] = p3

    # Структура каналов выручки год 3
    fig, ax = plt.subplots(figsize=(5.5, 4))
    rev3 = total_rev_k(FIN_Y3)
    b2b_pct = FIN_Y3["b2b_k"] / rev3 * 100
    b2c_pct = FIN_Y3["b2c_k"] / rev3 * 100
    api_pct = FIN_Y3["api_k"] / rev3 * 100
    labels = ["B2B клиники", "B2C Protocol", "B2B API/МИС"]
    sizes = [b2b_pct, b2c_pct, api_pct]
    ax.pie(sizes, labels=labels, autopct="%1.0f%%", colors=colors[:3], startangle=140, textprops={"fontsize": 8})
    ax.set_title(f"Структура выручки 2029 · SOM 8% TAM ({rev3} тыс. BYN/год)")
    fig.tight_layout()
    p4 = assets_dir / "chart_channels.png"
    fig.savefig(p4, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["channels"] = p4

    # EBITDA по годам
    fig, ax = plt.subplots(figsize=(7, 3.5))
    years = ["2027", "2028", "2029"]
    ebitda_vals = [ebitda_k(FIN_Y1), ebitda_k(FIN_Y2), ebitda_k(FIN_Y3)]
    bar_colors = [COLOR_NEGATIVE if v < 0 else colors[0] for v in ebitda_vals]
    bars = ax.bar(years, ebitda_vals, color=bar_colors)
    ax.axhline(0, color="#666", linewidth=0.8)
    ax.set_ylabel("тыс. BYN")
    ax.set_title(
        f"EBITDA год/мес · осторожный сценарий (8% TAM = {FIN_Y3['kz_month']:,} КЗ/мес)".replace(",", " ")
    )
    for bar, val in zip(bars, ebitda_vals):
        label = f"{val:+d}" if val != 0 else "0"
        y = val + (30 if val >= 0 else -40)
        ax.text(bar.get_x() + bar.get_width() / 2, y, label, ha="center", fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p5 = assets_dir / "chart_ebitda.png"
    fig.savefig(p5, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["ebitda"] = p5

    # Доля рынка и КЗ/мес
    fig, ax1 = plt.subplots(figsize=(7, 3.8))
    shares = [1, Y2_MARKET_SHARE * 100, Y3_MARKET_SHARE * 100]
    kz = [FIN_Y1["kz_month"] / 1000, FIN_Y2["kz_month"] / 1000, FIN_Y3["kz_month"] / 1000]
    x = range(3)
    ax1.bar([i - 0.2 for i in x], shares, width=0.35, label="Доля рынка, %", color=colors[3])
    ax2 = ax1.twinx()
    ax2.plot(list(x), kz, "o-", color=colors[4], linewidth=2, markersize=8, label="КЗ/мес, тыс.")
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(years)
    ax1.set_ylabel("Доля рынка, %")
    ax2.set_ylabel("КЗ/мес, тыс.")
    ax1.set_title(f"Доля от TAM 2,5 млн КЗ/мес · не весь рынок в год 3")
    lines1, labels1 = ax1.get_legend_handles_labels()
    lines2, labels2 = ax2.get_legend_handles_labels()
    ax1.legend(lines1 + lines2, labels1 + labels2, loc="upper left", fontsize=8)
    ax1.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p6 = assets_dir / "chart_market_share.png"
    fig.savefig(p6, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["market_share"] = p6

    # OPEX breakdown year 3 (stacked concept - grouped bars per year)
    fig, ax = plt.subplots(figsize=(7, 3.8))
    categories = ["ФОТ", "Инфра", "Маркетинг", "Прочее"]
    y1 = [180, 35, 35, 30]
    y2 = [270, 55, 55, 40]
    y3 = [420, 85, 85, 60]
    w = 0.25
    x = range(len(categories))
    for i, (vals, label, col) in enumerate(
        zip([y1, y2, y3], years, colors[:3])
    ):
        ax.bar([xi + (i - 1) * w for xi in x], vals, width=w, label=label, color=col)
    ax.set_xticks(list(x))
    ax.set_xticklabels(categories)
    ax.set_ylabel("тыс. BYN")
    ax.set_title("Структура OPEX по статьям (тыс. BYN)")
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p7 = assets_dir / "chart_opex.png"
    fig.savefig(p7, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["opex"] = p7

    # Маржа по тарифам B2B
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    tiers = ["0,99\nСтарт", "0,79\nКлиника", "0,69\nСеть"]
    prices = [0.99, 0.79, 0.69]
    cost_mid = 0.09
    margins = [(p - cost_mid) / p * 100 for p in prices]
    bars = ax.bar(tiers, margins, color=colors[:3])
    ax.set_ylabel("Валовая маржа L0, %")
    ax.set_title(f"Маржа L0 при себестоимости ~{cost_mid} BYN/КЗ")
    ax.set_ylim(0, 100)
    for bar, val in zip(bars, margins):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 1, f"{val:.0f}%", ha="center", fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p8 = assets_dir / "chart_margin.png"
    fig.savefig(p8, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["margin"] = p8

    # ROI якорного клиента
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    roi_labels = ["Затраты\nProtocol", "Экономия\nметодист", "Экономия\nЦИСЗ", "Итого\nэкономия"]
    roi_vals = [ROI_PROTOCOL_COST / 1000, ROI_METHODIST_SAVING / 1000, (ROI_TOTAL_SAVING - ROI_METHODIST_SAVING) / 1000, ROI_TOTAL_SAVING / 1000]
    bar_c = [COLOR_NEGATIVE, colors[0], colors[1], colors[2]]
    bars = ax.bar(roi_labels, roi_vals, color=bar_c)
    ax.set_ylabel("BYN/мес, тыс.")
    ax.set_title(f"ROI якорного клиента (нетто {ROI_NET / 1000:+.1f} тыс. BYN/мес)")
    for bar, val in zip(bars, roi_vals):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2, f"{val:.1f}", ha="center", fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p9 = assets_dir / "chart_roi.png"
    fig.savefig(p9, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["roi"] = p9

    # B2C воронка — согласована с FIN_Y3 (69,6 тыс. оплат в осторожном сценарии)
    fig, ax = plt.subplots(figsize=(7, 3.2))
    funnel = ["КЗ/год\n30 млн TAM", "Увидели QR/SMS\n2%", "Landing\n40%", "Оплата\n29%"]
    funnel_vals = [30_000_000, 600_000, 240_000, FIN_Y3["b2c_checks"]]
    display = [100, 2, 0.8, funnel_vals[3] / 300_000 * 100]
    ax.barh(funnel, display, color=[colors[2], colors[1], colors[4], colors[0]])
    ax.set_xlabel("Условная шкала (%)")
    ax.set_title(
        f"B2C-воронка → {FIN_Y3['b2c_checks']:,} платных/год (0,23% TAM, не 100%)".replace(",", " ")
    )
    for i, (bar, raw) in enumerate(zip(ax.patches, funnel_vals)):
        ax.text(bar.get_width() + 0.02, bar.get_y() + bar.get_height() / 2, f"{raw:,}".replace(",", " "), va="center", fontsize=8)
    ax.set_xlim(0, 110)
    fig.tight_layout()
    p10 = assets_dir / "chart_b2c_funnel.png"
    fig.savefig(p10, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["b2c_funnel"] = p10

    # B2B split: Кравира vs другие клиники
    fig, ax = plt.subplots(figsize=(7, 3.8))
    years = ["2027", "2028", "2029"]
    krav = [FIN_Y1["b2b_kravira_k"], FIN_Y2["b2b_kravira_k"], FIN_Y3["b2b_kravira_k"]]
    other = [FIN_Y1["b2b_other_k"], FIN_Y2["b2b_other_k"], FIN_Y3["b2b_other_k"]]
    x = range(len(years))
    ax.bar(x, krav, label="Кравира (якорь)", color=colors[0])
    ax.bar(x, other, bottom=krav, label="Другие ОЗ РБ", color=colors[4])
    ax.set_xticks(list(x))
    ax.set_xticklabels(years)
    ax.set_ylabel("тыс. BYN / год")
    ax.set_title("B2B: якорь Кравира vs другие ОЗ (осторожный план)")
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p11 = assets_dir / "chart_b2b_split.png"
    fig.savefig(p11, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["b2b_split"] = p11

    # B2C tier prices
    fig, ax = plt.subplots(figsize=(7, 3.8))
    tier_names = [t["name"].replace(" ", "\n") for t in B2C_TIERS]
    tier_prices = [t["price"] for t in B2C_TIERS]
    bars = ax.bar(tier_names, tier_prices, color=colors[: len(B2C_TIERS)])
    ax.axhline(B2C_AVG_PRICE, color=colors[5], linestyle="--", linewidth=1, label=f"Средний {B2C_AVG_PRICE} BYN")
    ax.set_ylabel("BYN за проверку")
    ax.set_title("Tier-цены B2C по specialty / сложности приёма")
    for bar, val in zip(bars, tier_prices):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 0.2, f"{val:.2f}", ha="center", fontsize=8)
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p12 = assets_dir / "chart_b2c_tiers.png"
    fig.savefig(p12, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["b2c_tiers"] = p12

    # B2C rev-share examples
    fig, ax = plt.subplots(figsize=(7, 3.5))
    examples = ["Промо\n2,99", "L2\n9,99", "Онко\n14,99", "Pre-op\n12,99"]
    ex_prices = [2.99, 9.99, 14.99, 12.99]
    clinic_shares = [clinic_revshare_byn(p)[0] for p in ex_prices]
    protocol_shares = [clinic_revshare_byn(p)[1] for p in ex_prices]
    x = range(len(examples))
    ax.bar(x, clinic_shares, label=f"Клинике {int(CLINIC_B2C_REVSHARE * 100)}%", color=colors[4])
    ax.bar(x, protocol_shares, bottom=clinic_shares, label="Protocol 70%", color=colors[0])
    ax.set_xticks(list(x))
    ax.set_xticklabels(examples)
    ax.set_ylabel("BYN с проверки")
    ax.set_title("Rev-share B2B2C: SMS/QR-ссылка от клиники")
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p13 = assets_dir / "chart_b2c_revshare.png"
    fig.savefig(p13, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["b2c_revshare"] = p13

    # B2C growth: три сценария Y3
    fig, ax = plt.subplots(figsize=(7, 3.8))
    scen_labels = ["2027", "2028", "2029\nостор.", "2029\nбазов.", "2029\nоптим."]
    scen_vals = [
        FIN_Y1["b2c_k"],
        FIN_Y2["b2c_k"],
        SCENARIO_CAUTIOUS["b2c_k"],
        SCENARIO_BASE["b2c_k"],
        SCENARIO_OPTIMISTIC["b2c_k"],
    ]
    bar_cols = [colors[0], colors[0], colors[0], colors[1], colors[4]]
    bars = ax.bar(scen_labels, scen_vals, color=bar_cols)
    ax.set_ylabel("тыс. BYN / год (Protocol, rev-share учтён)")
    ax.set_title(f"B2C: ~{B2C_PROTOCOL_PER_CHECK} BYN/проверка Protocol (микс tier + 30% клинике)")
    for bar, val in zip(bars, scen_vals):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 8, str(val), ha="center", fontsize=9)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p14 = assets_dir / "chart_b2c_growth.png"
    fig.savefig(p14, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["b2c_growth"] = p14

    # TAM → выручка (мост): показывает почему EBITDA << TAM
    fig, ax = plt.subplots(figsize=(7.5, 4))
    bridge_labels = ["TAM\nB2B теор.", "SAM\n5%", "SOM 8%\nB2B план", "B2C+API\n2029", "Выручка\nProtocol"]
    bridge_rev = [
        TAM_B2B_CEILING_YEAR_K,
        int(SAM_KZ_YEAR * 0.75 / 1000),
        FIN_Y3["b2b_k"],
        FIN_Y3["b2c_k"] + FIN_Y3["api_k"],
        total_rev_k(FIN_Y3),
    ]
    bar_c = [colors[2], colors[1], colors[0], colors[4], colors[3]]
    bars = ax.bar(bridge_labels, bridge_rev, color=bar_c)
    ax.set_ylabel("тыс. BYN / год")
    ax.set_title("TAM 2,5 млн КЗ/мес ≠ выручка: захват 8% B2B + 0,23% B2C")
    for bar, val in zip(bars, bridge_rev):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 200, f"{val:,}".replace(",", " "), ha="center", fontsize=7)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p15 = assets_dir / "chart_tam_bridge.png"
    fig.savefig(p15, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["tam_bridge"] = p15

    # EBITDA по сценариям Y3
    fig, ax = plt.subplots(figsize=(7, 3.8))
    scen_names = [s["label"].split("(")[0].strip() for s in ALL_SCENARIOS_Y3]
    ebitda_scen = [s["ebitda_k"] for s in ALL_SCENARIOS_Y3]
    ebitda_mo = [s["ebitda_month_k"] for s in ALL_SCENARIOS_Y3]
    bars = ax.bar(scen_names, ebitda_scen, color=[colors[0], colors[1], colors[4]])
    ax.set_ylabel("тыс. BYN / год")
    ax.set_title("EBITDA 2029: три сценария (подпись: год и ~месяц)")
    for bar, val, mo in zip(bars, ebitda_scen, ebitda_mo):
        ax.text(bar.get_x() + bar.get_width() / 2, bar.get_height() + 30, f"+{val}\n(~{mo}/мес)", ha="center", fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p16 = assets_dir / "chart_scenarios_ebitda.png"
    fig.savefig(p16, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["scenarios_ebitda"] = p16

    # Чувствительность EBITDA к доле рынка B2B
    fig, ax = plt.subplots(figsize=(7, 3.8))
    pen_labels = [f"{p[0]}%" for p in PENETRATION_SENSITIVITY]
    pen_vals = [p[1] for p in PENETRATION_SENSITIVITY]
    bar_cols = [colors[0] if p[0] == 8 else (colors[1] if p[0] == 10 else colors[2]) for p in PENETRATION_SENSITIVITY]
    bars = ax.bar(pen_labels, pen_vals, color=bar_cols)
    ax.axhline(ebitda_k(FIN_Y3), color=colors[5], linestyle="--", linewidth=1, label=f"План 8% = {ebitda_k(FIN_Y3)}")
    ax.set_xlabel("Доля TAM B2B (2,5 млн КЗ/мес)")
    ax.set_ylabel("EBITDA, тыс. BYN/год")
    ax.set_title("Чувствительность EBITDA к проникновению B2B (B2C фикс.)")
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p17 = assets_dir / "chart_penetration.png"
    fig.savefig(p17, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["penetration"] = p17

    # Вероятность успеха каналов
    fig, ax = plt.subplots(figsize=(7, 3.8))
    ch_names = [c["channel"] for c in CHANNEL_OUTLOOK]
    ch_prob = [c["prob"] * 100 for c in CHANNEL_OUTLOOK]
    ch_rev = [c["y3_k"] for c in CHANNEL_OUTLOOK]
    y_pos = range(len(ch_names))
    ax.barh(y_pos, ch_prob, color=colors[0], alpha=0.85)
    ax.set_yticks(list(y_pos))
    ax.set_yticklabels(ch_names, fontsize=8)
    ax.set_xlabel("Вероятность достижения плана к 2029, %")
    ax.set_title("Прогноз каналов: что с большей вероятностью «выстрелит»")
    for i, (p, r) in enumerate(zip(ch_prob, ch_rev)):
        ax.text(p + 1, i, f"{r} тыс. BYN", va="center", fontsize=7)
    ax.set_xlim(0, 105)
    ax.grid(axis="x", alpha=0.3)
    fig.tight_layout()
    p18 = assets_dir / "chart_channel_outlook.png"
    fig.savefig(p18, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["channel_outlook"] = p18

    # Смешанная выручка по сценариям Y3
    fig, ax = plt.subplots(figsize=(7, 4))
    x = range(len(ALL_SCENARIOS_Y3))
    b2b_s = [s["b2b_k"] for s in ALL_SCENARIOS_Y3]
    b2c_s = [s["b2c_k"] for s in ALL_SCENARIOS_Y3]
    api_s = [s["api_k"] for s in ALL_SCENARIOS_Y3]
    labels_s = ["Осторожн.", "Базовый", "Оптимист."]
    ax.bar(x, b2b_s, label="B2B", color=colors[0])
    ax.bar(x, b2c_s, bottom=b2b_s, label="B2C", color=colors[4])
    ax.bar(x, api_s, bottom=[b + c for b, c in zip(b2b_s, b2c_s)], label="API", color=colors[3])
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels_s)
    ax.set_ylabel("тыс. BYN / год")
    ax.set_title("Структура выручки 2029 по сценариям")
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p19 = assets_dir / "chart_scenarios_revenue.png"
    fig.savefig(p19, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["scenarios_revenue"] = p19

    # EBITDA помесячно 2027-2029
    fig, ax = plt.subplots(figsize=(7, 3.5))
    months = ["2027", "2028", "2029"]
    ebitda_annual = [ebitda_k(FIN_Y1), ebitda_k(FIN_Y2), ebitda_k(FIN_Y3)]
    ebitda_monthly = [round(v / 12, 1) for v in ebitda_annual]
    w = 0.35
    x = range(3)
    ax.bar([i - w / 2 for i in x], ebitda_annual, width=w, label="EBITDA год, тыс.", color=colors[0])
    ax.bar([i + w / 2 for i in x], ebitda_monthly, width=w, label="EBITDA мес, тыс.", color=colors[4])
    ax.axhline(0, color="#666", linewidth=0.8)
    ax.set_xticks(list(x))
    ax.set_xticklabels(months)
    ax.set_ylabel("тыс. BYN")
    ax.set_title(f"EBITDA год vs месяц · 2029: {ebitda_month_k(FIN_Y3)} тыс./мес при 8% TAM")
    ax.legend(fontsize=8)
    ax.grid(axis="y", alpha=0.3)
    fig.tight_layout()
    p20 = assets_dir / "chart_ebitda_monthly.png"
    fig.savefig(p20, dpi=170, bbox_inches="tight", facecolor=BG_FIG, edgecolor="none")
    plt.close(fig)
    paths["ebitda_monthly"] = p20

    return paths


def enrich_business_plan(doc: Document, charts: dict[str, Path]) -> None:
    """Таблицы в разделах и приложения перед блоком подписи."""
    # Раздел 6 - рынок
    p6 = find_section_body(doc, "6")
    if p6:
        anchor = add_table_after(
            p6,
            ["Показатель", "Значение", "Источник / допущение"],
            [
                ["КЗ/мес в МЦ «Кравира»", f"{KRAVIRA_KZ_MONTH:,}".replace(",", " "), "данные участника"],
                ["Доля Кравиры на рынке платных КЗ частных ОЗ РБ", "1%", "оценка участника"],
                ["КЗ/мес, сегмент частных ОЗ РБ", f"{MARKET_KZ_MONTH:,}".replace(",", " "), "25 000 / 0,01"],
                ["КЗ/год, сегмент", f"{MARKET_KZ_YEAR:,}".replace(",", " "), "× 12"],
                ["TAM @ 0,99 BYN/КЗ", f"~{TAM_REVENUE_YEAR:,} BYN/год".replace(",", " "), "30 млн × 0,99"],
                ["Цель год 3 (8% рынка)", f"{FIN_Y3['kz_month']:,} КЗ/мес".replace(",", " "), f"{Y3_MARKET_SHARE:.0%} TAM"],
                ["Цель год 2 (3% рынка)", f"{FIN_Y2['kz_month']:,} КЗ/мес".replace(",", " "), f"{Y2_MARKET_SHARE:.0%} TAM"],
            ],
            caption="Таблица 1. Оценка объёма рынка консультативных заключений",
        )
        anchor = add_picture_after(anchor, charts["market"])

        anchor = add_table_after(
            anchor,
            ["Показатель рынка РБ", "Значение", "Источник"],
            RB_MARKET_TABLE,
            caption="Таблица 1а. Контекст рынка платных медуслуг Республики Беларусь",
        )
        anchor = add_table_after(
            anchor,
            ["Драйвер спроса на Protocol", "Пояснение"],
            list(CISZ_DRIVERS),
            caption="Таблица 9. Факторы роста спроса",
        )
        add_picture_after(anchor, charts["market_share"])

    # Раздел 8 - маркетинг
    p8sec = find_section_body(doc, "8")
    if p8sec:
        add_table_after(
            p8sec,
            ["Канал", "Целевая аудитория", "KPI", "Бюджет/год"],
            [
                ["Прямые продажи B2B", "частные ОЗ >5k КЗ/мес", "3 договора/год", "40 тыс. BYN"],
                ["OEM API Айболит", "все клиенты МИС", "1 интеграция", "в CAPEX"],
                ["QR B2C в Кравире", "пациенты", "0,1% конверсия", "15 тыс. BYN"],
                ["Конференции ЦИСЗ", "методслужбы", "2 выступления", "10 тыс. BYN"],
                ["Контент SEO", "физлица", "5k визитов/мес", "20 тыс. BYN"],
            ],
            caption="Таблица 1б. Маркетинговый план",
        )

    # Раздел 9 - цены
    p9 = find_section_body(doc, "9")
    if p9:
        anchor = add_table_after(
            p9,
            ["Тариф B2B", "Цена L0", "Условия"],
            [
                ["Старт", "0,99 BYN/КЗ", "до 1 000 КЗ/мес"],
                ["Клиника", "0,79 BYN/КЗ", "до 10 000 КЗ/мес, мин. 5 000 BYN/мес"],
                ["Сеть", "0,69 BYN/КЗ", "25 000+ КЗ/мес, мин. 12 000 BYN/мес"],
                ["L2 методист", "+0,50 BYN/КЗ", "углублённый разбор"],
                ["Внедрение API", "15-40 тыс. BYN", "разово + поддержка"],
            ],
            caption="Таблица 2. Прайс-лист B2B",
        )
        anchor = add_table_after(
            anchor,
            ["Тариф B2C", "Цена", "Содержание"],
            [
                ["Базовая проверка L1", "4,99 BYN", "структурный разбор без LLM"],
                ["Подробный отчёт L2", "9,99 BYN", "пояснения простым языком"],
                ["Абонемент", "12,99 BYN/мес", "3 проверки"],
                ["Промо от клиники", "2,99 BYN", "QR на чеке/КЗ"],
            ],
            caption="Таблица 3. Прайс-лист B2C (физические лица)",
        )
        anchor = add_table_after(
            anchor,
            ["Статья себестоимости L0", "BYN/КЗ"],
            [
                ["Сервер / амортизация", "0,02-0,05"],
                ["Электричество, админ", "0,01"],
                ["Поддержка (0,5 FTE на 100k КЗ)", "0,03-0,06"],
                ["Итого L0", "0,06-0,12"],
                ["Маржа при 0,99 BYN", "~85-90%"],
            ],
            caption="Таблица 4. Себестоимость проверки L0",
        )
        anchor = add_table_after(
            anchor,
            ["Тариф", "Цена", "Себест.", "Маржа L0"],
            [
                ["Старт", "0,99", "0,09", "~91%"],
                ["Клиника", "0,79", "0,09", "~89%"],
                ["Сеть", "0,69", "0,09", "~87%"],
            ],
            caption="Таблица 9. Маржа по тарифам B2B",
        )
        anchor = add_picture_after(anchor, charts["pricing"])
        add_picture_after(anchor, charts["margin"])

    # Раздел 10 - конкуренты
    p10 = find_section_body(doc, "10")
    if p10:
        anchor = add_table_after(
            p10,
            ["Альтернатива", "Охват", "Слабость", "Стоимость"],
            COMPETITOR_MATRIX,
            caption="Таблица 5. Конкурентный анализ (детально)",
        )
        add_table_after(
            anchor,
            ["Критерий", "Protocol", "Ручной аудит", "LLM-чат", "Шаблон МИС"],
            [
                ["КП Минздрава РБ", "5", "4", "1", "2"],
                ["100% потока", "5", "1", "2", "5"],
                ["ЦИСЗ / FHIR BY", "5", "3", "0", "3"],
                ["Evidence map", "5", "4", "1", "1"],
                ["B2C для пациента", "5", "0", "2", "0"],
            ],
            caption="Таблица 5а. Оценка по 5-балльной шкале",
        )

    # Раздел 13 - команда
    p13 = find_section_body(doc, "13")
    if p13:
        add_table_after(
            p13,
            ["Роль", "FTE", "Функция"],
            [
                ["Руководитель проекта / главврач", "0,2", "методология, Минздрав"],
                ["Ведущий разработчик", "1,0", "продукт, API, МИС"],
                ["Методист-клиницист", "0,5", "правила, валидация"],
                ["Менеджер B2B/B2C", "0,5", "продажи, договоры"],
                ["ИТ-админ", "0,2", "on-prem, ИБ"],
            ],
            caption="Таблица 6. Организационная структура проекта",
        )

    # Раздел 14 - риски
    p14 = find_section_body(doc, "14")
    if p14:
        add_table_after(
            p14,
            ["Риск", "Вероятность", "Митигация"],
            [
                ["ПДн B2C", "средняя", "минимизация хранения, HTTPS, on-prem L0"],
                ["Изменение КП Минздрава", "высокая", "автосинхронизация корпуса"],
                ["Низкая готовность платить", "средняя", "ROI: меньше отказов ЦИСЗ"],
                ["Регуляторика «медизделие»", "низкая", "позиция СПО/ИС"],
                ["Низкая конверсия B2C", "средняя", "QR-кампании в клиниках"],
            ],
            caption="Таблица 7. Матрица рисков",
        )

    # Раздел 15 - финансы
    p15 = find_section_body(doc, "15")
    if p15:
        anchor = add_table_after(
            p15,
            ["Показатель", "2027", "2028", "2029"],
            [
                ["Клиенты (ОЗ экв.)", "1 (Кравира)", "3", "8"],
                ["КЗ/мес, всего", f"{FIN_Y1['kz_month']:,}".replace(",", " "), f"{FIN_Y2['kz_month']:,}".replace(",", " "), f"{FIN_Y3['kz_month']:,}".replace(",", " ")],
                ["Доля рынка РБ", "1%", "3%", "8%"],
                ["Выручка B2B, тыс. BYN", str(FIN_Y1["b2b_k"]), str(FIN_Y2["b2b_k"]), str(FIN_Y3["b2b_k"])],
                ["Выручка B2C, тыс. BYN", str(FIN_Y1["b2c_k"]), str(FIN_Y2["b2c_k"]), str(FIN_Y3["b2c_k"])],
                ["Выручка API/МИС, тыс. BYN", str(FIN_Y1["api_k"]), str(FIN_Y2["api_k"]), str(FIN_Y3["api_k"])],
                ["Выручка итого, тыс. BYN", str(FIN_Y1["b2b_k"] + FIN_Y1["b2c_k"] + FIN_Y1["api_k"]), str(FIN_Y2["b2b_k"] + FIN_Y2["b2c_k"] + FIN_Y2["api_k"]), str(FIN_Y3["b2b_k"] + FIN_Y3["b2c_k"] + FIN_Y3["api_k"])],
                ["OPEX, тыс. BYN", str(FIN_Y1["opex_k"]), str(FIN_Y2["opex_k"]), str(FIN_Y3["opex_k"])],
                ["EBITDA, тыс. BYN", str(ebitda_k(FIN_Y1)), f"+{ebitda_k(FIN_Y2)}", f"+{ebitda_k(FIN_Y3)}"],
            ],
            caption="Таблица 8. Финансовый план на 3 года (тыс. BYN)",
        )
        anchor = add_table_after(
            anchor,
            ["Статья OPEX", "2027", "2028", "2029"],
            [
                ["ФОТ", "180", "270", "420"],
                ["Инфраструктура", "35", "55", "85"],
                ["Маркетинг", "35", "55", "85"],
                ["Прочее", "30", "40", "60"],
                ["Итого", "280", "420", "650"],
            ],
            caption="Таблица 11. Структура OPEX (тыс. BYN)",
        )
        anchor = add_table_after(
            anchor,
            ["Статья инвестиций", "Сумма, BYN", "Срок"],
            INVESTMENT_PLAN,
            caption="Таблица 10. План инвестиций 2026-2027",
        )
        anchor = add_table_after(
            anchor,
            ["Показатель", "2027", "2028", "2029"],
            [
                ["EBITDA, тыс. BYN", str(ebitda_k(FIN_Y1)), f"+{ebitda_k(FIN_Y2)}", f"+{ebitda_k(FIN_Y3)}"],
                ["Накопленный CF, тыс. BYN", str(ebitda_k(FIN_Y1)), str(ebitda_k(FIN_Y1) + ebitda_k(FIN_Y2)), str(ebitda_k(FIN_Y1) + ebitda_k(FIN_Y2) + ebitda_k(FIN_Y3))],
                ["Сертификат ГКНТ", f"{CERTIFICATE_BYN:,} BYN".replace(",", " "), "-", "-"],
            ],
            caption="Таблица 12. Cash flow (упрощённо, тыс. BYN)",
        )
        anchor = add_picture_after(anchor, charts["revenue"])
        anchor = add_picture_after(anchor, charts["ebitda"])
        add_picture_after(anchor, charts["opex"])

    # Раздел 16 - иные сведения + ссылка на приложения
    p16 = find_section_body(doc, "16")
    if p16:
        p16.clear()
        run = p16.add_run(
            dash(
                "Бизнес-план соответствует рекомендуемому объёму 10-40 страниц: основной текст - "
                "разделы 1-16; детальные расчёты, таблицы и графики - в приложениях А-Е и в файле "
                "05_Biznes_plan_Prilozheniya.html (для печати в PDF). "
                "Дополнительно прилагаются architecture-kravira-fhir-mis.pdf и mvp-presentation.html."
            )
        )
        run.font.size = Pt(11)

    # Приложения перед подписью
    sig_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Участник конкурса"):
            sig_idx = i
            break
    if sig_idx is None:
        return

    sig_p = doc.paragraphs[sig_idx - 1] if sig_idx > 0 else doc.paragraphs[sig_idx]
    anchor = sig_p
    if sig_idx > 0 and not doc.paragraphs[sig_idx - 1].text.strip():
        anchor = doc.paragraphs[sig_idx - 1]

    anchor = add_paragraph_after(anchor, "", space_before=12)
    anchor = add_paragraph_after(anchor, "ПРИЛОЖЕНИЯ К БИЗНЕС-ПЛАНУ", bold=True, size=14, space_after=10)

    appendix_sections: list[tuple[str, str, list[tuple[str, ...]] | None, str | None]] = [
        (
            "Приложение А. Сводный прайс-лист B2B/B2C",
            "Коммерческие тарифы для клиник и физических лиц.",
            [
                ("Канал", "Продукт", "Цена", "Примечание"),
                ("B2B", "L0 Старт", "0,99 BYN/КЗ", "pay-as-you-go"),
                ("B2B", "L0 Клиника", "0,79 BYN/КЗ", "от 10k КЗ/мес"),
                ("B2B", "L0 Сеть", "0,69 BYN/КЗ", "Кравира, 25k+"),
                ("B2B", "API МИС", "индивид.", "15-40k внедрение"),
                ("B2C", "L1 Базовый", "4,99 BYN", "терапия"),
                ("B2C", "L1+ Расширенный", "6,99 BYN", "специалист+анализы"),
                ("B2C", "L2 Стандарт", "9,99 BYN", "подробный разбор"),
                ("B2C", "L2 Онкология", "14,99 BYN", "онко, ЗНО"),
                ("B2C", "L2 Pre-op", "12,99 BYN", "предоперационный"),
                ("B2C", "Промо SMS/QR", "2,99 BYN", "rev-share 30% клинике"),
            ],
            None,
        ),
        (
            "Приложение Б. Дорожная карта 2025-2029",
            "Ключевые вехи коммерциализации Protocol.",
            [
                ("Этап", "Срок", "Результат"),
                ("Пилот L0", "2025-Q2 2026", "метрики в Кравире"),
                ("API МИС Айболит", "Q3 2026-Q1 2027", "L0 при сохранении КЗ"),
                ("B2C beta", "Q4 2026", "витрина 4,99/9,99 BYN"),
                ("3-5 B2B клиентов", "2027", "договоры, пакеты"),
                ("5-15% рынка", "2028-2029", "масштабирование"),
            ],
            None,
        ),
        (
            "Приложение В. KPI пилота (целевые)",
            "Показатели для подтверждения эффективности на площадке Кравиры.",
            [
                ("KPI", "Целевое значение", "Как измерять"),
                ("Время L0", "< 2 с", "лог API"),
                ("Покрытие потока", "100% КЗ", "интеграция МИС"),
                ("gate_score < 70", "снижение vs база", "отчёт Protocol"),
                ("Критические пробелы ЦИСЗ", "снижение", "cisz_readiness"),
                ("Доработки после ЭЦП", "-30%", "сравнение периодов"),
            ],
            None,
        ),
        (
            "Приложение Г. График выручки B2B/B2C",
            "Осторожный сценарий на 3 года.",
            None,
            "revenue",
        ),
        (
            "Приложение Д. Структура каналов монетизации (год 3)",
            "Доля выручки по каналам.",
            None,
            "channels",
        ),
        (
            "Приложение Е. EBITDA и рост доли рынка",
            "Финансовая динамика и масштабирование.",
            None,
            "ebitda",
        ),
        (
            "Приложение Ж. ROI якорного клиента (МЦ «Кравира»)",
            "Сравнение затрат Protocol и экономии методиста + ЦИСЗ.",
            None,
            "roi",
        ),
        (
            "Приложение З. B2C-воронка и OPEX",
            "Консервативная модель конверсии физлиц и структура расходов.",
            None,
            "b2c_funnel",
        ),
    ]

    for title, intro, table_data, chart_key in appendix_sections:
        anchor = add_paragraph_after(anchor, title, bold=True, size=12, space_before=14, space_after=4)
        anchor = add_paragraph_after(anchor, intro, size=10, space_after=6)
        if table_data:
            headers = table_data[0]
            rows = table_data[1:]
            anchor = add_table_after(anchor, headers, rows)
        if chart_key and chart_key in charts:
            anchor = add_picture_after(anchor, charts[chart_key])
